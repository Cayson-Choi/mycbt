"""
시험 문제를 Word(.docx) 파일로 내보내기.

- KaTeX/LaTeX 수식 → Pandoc으로 Word 네이티브 수식(OMML) 변환
- Cloudinary 이미지 자동 다운로드 후 임베드
- A4 2단 레이아웃 + 머릿말(시험명) + 문제 사이 가는 회색 구분선
- 본문 10pt, 분수 안 텍스트 15pt
- 이미지 최대 7×5cm 비율 유지 축소

DB 접속 정보는 ``.env.local``의 ``DATABASE_URL``에서 읽습니다.

사용법::

    python scripts/export-exam-to-docx.py --list
    python scripts/export-exam-to-docx.py --exam-id 290
    python scripts/export-exam-to-docx.py --exam-id 290 --no-answers --no-images
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from contextlib import closing
from pathlib import Path

import psycopg2
import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu

sys.stdout.reconfigure(encoding='utf-8')


# ─── 상수 ────────────────────────────────────────────────────────────

CIRCLED = {1: '①', 2: '②', 3: '③', 4: '④'}

# OMML(수식) namespace
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
def mtag(tag: str) -> str:
    return f'{{{M_NS}}}{tag}'

# 글자 크기 (half-points) — Word는 sz를 half-points로 표현
BODY_SZ = '20'   # 10pt
FRAC_SZ = '30'   # 15pt

# 이미지 최대 크기 (비율 유지하며 작은 쪽 비율로 축소)
MAX_IMG_W = Cm(7.0)
MAX_IMG_H = Cm(5.0)

# 컬럼 사이 간격 (twips, 425 ≈ 0.75cm)
COLUMN_SPACE = '425'

# 문제 사이 구분선 (가는 회색)
SEPARATOR_COLOR = 'BBBBBB'
SEPARATOR_SZ = '4'  # 0.5pt


# ─── DB ─────────────────────────────────────────────────────────────

def load_database_url() -> str:
    """`.env.local` 또는 환경변수에서 DATABASE_URL을 읽는다."""
    if (url := os.environ.get('DATABASE_URL')):
        return url.strip().strip('"').strip("'")

    env_path = Path(__file__).resolve().parent.parent / '.env.local'
    if env_path.exists():
        for line in env_path.read_text(encoding='utf-8').splitlines():
            line = line.strip()
            if line.startswith('DATABASE_URL='):
                value = line.split('=', 1)[1].strip().strip('"').strip("'")
                return value
    raise RuntimeError('DATABASE_URL 환경변수 또는 .env.local에 DATABASE_URL이 필요합니다.')


def list_exams(conn) -> None:
    with closing(conn.cursor()) as cur:
        cur.execute("""
            SELECT
                e.id, c.name, e.name, e.year, e.round, e.exam_type,
                COUNT(q.id) AS qcount
            FROM exams e
            JOIN exam_categories c ON c.id = e.category_id
            LEFT JOIN questions q ON q.exam_id = e.id AND q.is_active = true
            GROUP BY e.id, c.name, e.name, e.year, e.round, e.exam_type
            HAVING COUNT(q.id) > 0
            ORDER BY c.name, e.year DESC NULLS LAST, e.round DESC NULLS LAST, e.id
        """)
        rows = cur.fetchall()

    print(f"\n{'ID':>4}  {'카테고리':20}  {'시험명':40}  {'년도':>5}  {'회차':>4}  {'유형':10}  {'문항':>5}")
    print('-' * 110)
    for eid, cat, name, year, rnd, etype, qcount in rows:
        print(f"{eid:>4}  {cat:20}  {name:40}  {str(year or ''):>5}  {str(rnd or ''):>4}  {etype:10}  {qcount:>5}")
    print(f"\n총 {len(rows)}개 시험\n")


def fetch_exam(conn, exam_id: int):
    with closing(conn.cursor()) as cur:
        cur.execute("""
            SELECT e.id, c.name, e.name, e.year, e.round, e.exam_type, e.duration_minutes
            FROM exams e
            JOIN exam_categories c ON c.id = e.category_id
            WHERE e.id = %s
        """, (exam_id,))
        row = cur.fetchone()
        if not row:
            return None, [], []
        exam = {
            'id': row[0], 'category': row[1], 'name': row[2],
            'year': row[3], 'round': row[4], 'exam_type': row[5],
            'duration': row[6],
        }

        cur.execute("""
            SELECT id, name, order_no, questions_per_attempt
            FROM subjects WHERE exam_id = %s
            ORDER BY order_no, id
        """, (exam_id,))
        subjects = [
            {'id': r[0], 'name': r[1], 'order_no': r[2], 'qpa': r[3]}
            for r in cur.fetchall()
        ]

        cur.execute("""
            SELECT id, subject_id, question_code, question_text,
                   choice_1, choice_2, choice_3, choice_4,
                   choice_1_image, choice_2_image, choice_3_image, choice_4_image,
                   answer, answer_text, explanation, image_url, points
            FROM questions
            WHERE exam_id = %s AND is_active = true
            ORDER BY subject_id, id
        """, (exam_id,))
        questions = [
            {
                'id': r[0], 'subject_id': r[1], 'code': r[2], 'q': r[3],
                'c1': r[4], 'c2': r[5], 'c3': r[6], 'c4': r[7],
                'c1_img': r[8], 'c2_img': r[9], 'c3_img': r[10], 'c4_img': r[11],
                'answer': r[12], 'answer_text': r[13], 'explanation': r[14],
                'image_url': r[15], 'points': r[16],
            }
            for r in cur.fetchall()
        ]
    return exam, subjects, questions


# ─── Markdown 빌더 ──────────────────────────────────────────────────

def normalize_math(text: str) -> str:
    """필요 시 KaTeX → LaTeX 변환을 위한 hook (현재는 통과)."""
    return text or ''


def build_exam_title(exam: dict) -> str:
    parts = []
    if exam.get('year'):
        parts.append(f"{exam['year']}년")
    if exam.get('round'):
        parts.append(f"{exam['round']}회")
    parts.append(exam['name'])
    return ' '.join(parts)


def _format_multiline_body(prefix: str, text: str) -> list[str]:
    """`prefix` + 본문을 같은 단락으로, 본문 안 줄바꿈은 markdown line break(`\\`)로."""
    body_lines = (text or '').split('\n')
    if len(body_lines) == 1:
        return [f"{prefix} {body_lines[0]}"]
    head = f"{prefix} {body_lines[0]}\\"
    tail = [line + '\\' for line in body_lines[1:-1]] + [body_lines[-1]]
    return [head, *tail]


def format_question_md(q: dict, q_num: int, *, include_answers: bool, include_images: bool, include_explanation: bool) -> str:
    lines: list[str] = []

    lines.extend(_format_multiline_body(f"**{q_num}.**", normalize_math(q['q'])))
    lines.append('')

    if include_images and q.get('image_url'):
        lines.append(f"![]({q['image_url']})")
        lines.append('')

    for n in (1, 2, 3, 4):
        ctext = normalize_math(q.get(f'c{n}') or '')
        cimg = q.get(f'c{n}_img')
        line = f"{CIRCLED[n]} {ctext}"
        if include_images and cimg:
            line += f" ![]({cimg})"
        lines.append(line)
        lines.append('')

    if include_answers:
        ans = q.get('answer')
        if ans and ans in CIRCLED:
            lines.append(f"**정답: {CIRCLED[ans]}**")
        elif ans == 0:
            lines.append("**정답: ⚠️ 정답불명 (관리자 확인 필요)**")
        elif q.get('answer_text'):
            lines.append(f"**정답:** {q['answer_text']}")
        else:
            lines.append("**정답: -**")
        lines.append('')

        if include_explanation and q.get('explanation'):
            lines.append("**해설:** " + normalize_math(q['explanation']).replace('\n', ' '))
            lines.append('')

    lines.append('---')
    lines.append('')
    return '\n'.join(lines)


def build_markdown(exam: dict, subjects: list, questions: list, *, include_answers: bool, include_images: bool, include_explanation: bool) -> str:
    by_subject: dict[int, list] = {}
    for q in questions:
        by_subject.setdefault(q['subject_id'], []).append(q)

    md: list[str] = []
    global_num = 0
    for s in sorted(subjects, key=lambda x: (x['order_no'], x['id'])):
        qs = by_subject.get(s['id'], [])
        if not qs:
            continue
        md.append(f"**{s['order_no']}과목. {s['name']}**")
        md.append('')
        for q in qs:
            global_num += 1
            md.append(format_question_md(
                q, global_num,
                include_answers=include_answers,
                include_images=include_images,
                include_explanation=include_explanation,
            ))
    return '\n'.join(md)


# ─── docx 후처리 ────────────────────────────────────────────────────

def _ensure_w_rPr(run, after_m_rPr: bool = False):
    """run(w:r 또는 m:r) 안에 w:rPr이 없으면 만들고 반환."""
    w_rPr = run.find(qn('w:rPr'))
    if w_rPr is not None:
        return w_rPr
    w_rPr = OxmlElement('w:rPr')
    if after_m_rPr:
        m_rPr = run.find(mtag('rPr'))
        if m_rPr is not None:
            m_rPr.addnext(w_rPr)
            return w_rPr
    run.insert(0, w_rPr)
    return w_rPr


def _force_size(w_rPr, val: str) -> None:
    for tag in ('w:sz', 'w:szCs'):
        el = w_rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            w_rPr.append(el)
        el.set(qn('w:val'), val)


def _set_section_layout(section, header_title: str | None) -> None:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    if header_title:
        p = section.header.paragraphs[0]
        p.text = ''
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_title)
        run.bold = True

    sectPr = section._sectPr
    for old in sectPr.findall(qn('w:cols')):
        sectPr.remove(old)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), COLUMN_SPACE)
    cols.set(qn('w:equalWidth'), '1')
    cols.set(qn('w:sep'), '1')
    sectPr.append(cols)


def _replace_vml_separators(doc) -> None:
    """Pandoc은 markdown `---`을 VML rectangle로 변환 → Word에서 두 줄로 보임.
    VML hr을 제거하고 paragraph bottom border(가는 회색 단일 선)로 교체.
    """
    for p in doc.paragraphs:
        p_xml = p._p
        picts = p_xml.findall('.//' + qn('w:pict'))
        has_hr = False
        for pict in picts:
            for child in pict:
                if child.tag.endswith('rect'):
                    has_hr = True
                    break
            r = pict.getparent()
            while r is not None and r.tag != qn('w:r'):
                r = r.getparent()
            if r is not None and (gp := r.getparent()) is not None:
                gp.remove(r)
        if not has_hr:
            continue

        pPr = p_xml.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_xml.insert(0, pPr)
        for old in pPr.findall(qn('w:pBdr')):
            pPr.remove(old)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), SEPARATOR_SZ)
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), SEPARATOR_COLOR)
        pBdr.append(bottom)
        pPr.append(pBdr)


def _force_font_sizes(doc) -> None:
    """본문 10pt, OMML 분수 안 15pt 강제.
    OMML inline math의 fraction이 자동 축소되는 문제를 우회.
    """
    # 1) 일반 텍스트 (w:r) → 10pt
    for p in doc.paragraphs:
        for r in p._p.iter(qn('w:r')):
            _force_size(_ensure_w_rPr(r), BODY_SZ)

    # 2) OMML 수식 텍스트 (m:r) → 10pt (분수 밖 영역)
    for p in doc.paragraphs:
        for r in p._p.iter(mtag('r')):
            _force_size(_ensure_w_rPr(r, after_m_rPr=True), BODY_SZ)

    # 3) 분수(m:f) 안 텍스트 → 15pt (덮어쓰기)
    for f_elem in doc.element.iter(mtag('f')):
        for r in f_elem.iter(mtag('r')):
            _force_size(_ensure_w_rPr(r, after_m_rPr=True), FRAC_SZ)


def _resize_images(doc) -> None:
    """이미지를 MAX_IMG_W × MAX_IMG_H 안에 비율 유지하며 축소."""
    for shape in doc.inline_shapes:
        try:
            w, h = shape.width, shape.height
            if not w or not h:
                continue
            scale = min(MAX_IMG_W / w, MAX_IMG_H / h, 1.0)
            if scale < 1.0:
                shape.width = Emu(int(w * scale))
                shape.height = Emu(int(h * scale))
        except Exception:
            pass


def apply_two_column_layout(docx_path: str, header_title: str | None = None) -> None:
    """Pandoc 변환 결과 docx를 시험지 형식으로 후처리."""
    doc = Document(docx_path)
    for section in doc.sections:
        _set_section_layout(section, header_title)
    _replace_vml_separators(doc)
    _force_font_sizes(doc)
    _resize_images(doc)
    doc.save(docx_path)


# ─── 출력 파일명 ────────────────────────────────────────────────────

def safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', '_', s)
    s = re.sub(r'\s+', '_', s).strip('_')
    return s


def default_out_path(exam: dict) -> str:
    parts = [
        str(exam.get('year') or ''),
        f"{exam.get('round')}회" if exam.get('round') else '',
        exam['name'],
    ]
    base = safe_filename('_'.join(p for p in parts if p))
    os.makedirs('data/exports', exist_ok=True)
    return os.path.join('data', 'exports', f'{base}.docx')


# ─── main ──────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument('--list', action='store_true', help='시험 목록 조회')
    parser.add_argument('--exam-id', type=int, help='내보낼 시험 ID')
    parser.add_argument('--no-answers', action='store_true', help='정답 제외')
    parser.add_argument('--no-images', action='store_true', help='이미지 제외')
    parser.add_argument('--with-explanation', action='store_true', help='해설 포함')
    parser.add_argument('--out', type=str, help='출력 파일 경로(.docx)')
    args = parser.parse_args()

    with closing(psycopg2.connect(load_database_url())) as conn:
        if args.list or not args.exam_id:
            list_exams(conn)
            if not args.exam_id:
                print('--exam-id <ID> 로 시험을 지정하세요.')
                return

        exam, subjects, questions = fetch_exam(conn, args.exam_id)

    if not exam:
        print(f'❌ exam id={args.exam_id} 를 찾을 수 없습니다.')
        return
    if not questions:
        print('❌ 활성 문제가 없습니다.')
        return

    print(f"📘 {exam['category']} / {exam['name']} ({exam.get('year')}년 {exam.get('round')}회)")
    print(f"   문항 수: {len(questions)}, 과목 수: {len(subjects)}")

    md = build_markdown(
        exam, subjects, questions,
        include_answers=not args.no_answers,
        include_images=not args.no_images,
        include_explanation=args.with_explanation,
    )

    out_path = args.out or default_out_path(exam)
    md_debug_path = out_path.replace('.docx', '.md')
    Path(md_debug_path).write_text(md, encoding='utf-8')
    print(f"📄 markdown: {md_debug_path}")

    print('⏳ Pandoc 변환 중...')
    pypandoc.convert_text(
        md,
        to='docx',
        format='markdown+tex_math_dollars+raw_tex',
        outputfile=out_path,
        extra_args=['--standalone', '--quiet'],
    )

    print('🪄 2단 레이아웃 적용 중...')
    apply_two_column_layout(out_path, header_title=build_exam_title(exam))

    print(f'✅ 저장 완료: {out_path}')


if __name__ == '__main__':
    main()
